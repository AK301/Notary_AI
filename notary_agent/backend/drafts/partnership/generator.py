from docx import Document
import os
from backend.drafts.partnership.schema import (
    PartnershipDeedRequest,
)  # Adjust path if needed


def generate_partnership_docx(
    data: PartnershipDeedRequest, filename_prefix: str = "partnership_deed"
) -> str:
    doc = Document()

    doc.add_heading("PARTNERSHIP DEED", level=1)

    doc.add_paragraph(
        f"THIS DEED OF PARTNERSHIP is executed on this {data.execution_date} at {data.jurisdiction} by and between:"
    )

    doc.add_paragraph(
        f"1. {data.partner1_full_name}, Son of {data.partner1_father_name}, Age {data.partner1_age}, residing at {data.partner1_address} (Hereinafter referred to as Partner No. 1)"
    )

    doc.add_paragraph("AND")

    doc.add_paragraph(
        f"2. {data.partner2_full_name}, Son of {data.partner2_father_name}, Age {data.partner2_age}, residing at {data.partner2_address} (Hereinafter referred to as Partner No. 2)"
    )

    doc.add_paragraph(
        "WHEREAS the above-named partners have mutually agreed to carry on the business in partnership under the terms and conditions set forth hereinbelow:"
    )

    clauses = [
        (
            "1. NAME OF THE FIRM:",
            f"The name of the partnership firm shall be {data.firm_name}.",
        ),
        (
            "2. NATURE OF BUSINESS:",
            f"The business to be carried on shall be {data.business_type}.",
        ),
        (
            "3. PLACE OF BUSINESS:",
            f"The place of business shall be {data.business_address}, and the area of operation will be {data.area_of_operation}.",
        ),
        (
            "4. DURATION:",
            f"The partnership shall commence on {data.start_date} and shall be a Partnership at Will.",
        ),
        (
            "5. CAPITAL CONTRIBUTION:",
            f"Partner No.1: ₹{data.capital_contribution.split(',')[0]}, Partner No.2: ₹{data.capital_contribution.split(',')[1]}",
        ),
        (
            "6. PROFIT AND LOSS SHARING:",
            f"Partner No.1: {data.profit_sharing.split(',')[0]}, Partner No.2: {data.profit_sharing.split(',')[1]}",
        ),
        ("7. DUTIES AND RESPONSIBILITIES:", data.duties),
        (
            "8. BANKING OPERATIONS:",
            "The firm shall open a bank account in its name and the same shall be operated jointly or individually as agreed by the partners.",
        ),
        (
            "9. ACCOUNTS AND AUDIT:",
            "Proper books of account shall be maintained and closed on 31st March every year.",
        ),
        (
            "10. ADMISSION OF NEW PARTNER:",
            "No new partner shall be admitted without the consent of all existing partners.",
        ),
        (
            "11. RETIREMENT AND DEATH:",
            "On retirement or death of any partner, the firm may be dissolved or continued as mutually agreed.",
        ),
        (
            "12. ARBITRATION:",
            "Any disputes between the partners shall be referred to arbitration in accordance with the Arbitration and Conciliation Act, 1996.",
        ),
        (
            "13. GOVERNING LAW:",
            "This deed shall be governed by and construed in accordance with the laws of India.",
        ),
    ]

    for clause_title, clause_content in clauses:
        p = doc.add_paragraph()
        p.add_run(clause_title).bold = True
        p.add_run("\n" + clause_content)

    doc.add_paragraph(
        "\nIN WITNESS WHEREOF, the parties hereto have signed this deed on the date and place mentioned above."
    )

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Partner No. 1 Signature:\n\n\n_______________________"
    table.cell(0, 1).text = "Partner No. 2 Signature:\n\n\n_______________________"
    table.cell(1, 0).text = data.partner1_full_name
    table.cell(1, 1).text = data.partner2_full_name

    output_dir = "generated_docs"
    os.makedirs(output_dir, exist_ok=True)

    safe_filename = f"{filename_prefix.replace(' ', '_')}.docx"
    file_path = os.path.join(output_dir, safe_filename)
    doc.save(file_path)

    return file_path
