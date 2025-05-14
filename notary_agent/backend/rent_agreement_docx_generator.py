from docx import Document
import os

def generate_rent_agreement_docx(data: dict, filename_prefix="rent_agreement"):
    doc = Document()

    doc.add_heading("RENT AGREEMENT", 0)

    doc.add_paragraph(
        f"This Rent Agreement is made on this {data['agreement_date']} by {data['landlord']['name']} S/o {data['landlord']['father_name']}, "
        f"residing at {data['landlord']['address']}, hereinafter referred to as the Lessor/Owner, Party of the First Part."
    )

    doc.add_paragraph(
        f"AND {data['tenant']['company_name']}, through its proposed director {data['tenant']['director_name']}, hereinafter referred to as the Lessee/Tenant, Party of the Second Part."
    )

    doc.add_paragraph(
        f"WHEREAS the Lessor/Owner is the owner and in possession of the property No: {data['property']['number']}, {data['property']['address']}, "
        f"and has agreed to let out one office room, one toilet & bathroom set on said property to the Lessee/Tenant, and the Lessee/Tenant has agreed to take the same on rent of Rs. {data['rent']['amount']}/- "
        f"({data['rent']['amount_words']}) per month."
    )

    doc.add_paragraph("\nNOW THIS RENT AGREEMENT WITNESSETH AS UNDER:")

    clauses = [
        f"That the Tenant/Lessee shall pay a monthly rent of Rs. {data['rent']['amount']}/- ({data['rent']['amount_words']}) per month, excluding electricity and water charges.",
        "That the Tenant / Lessee shall not sub-let any part of the above said premises to anyone else without the consent of the Owner.",
        "That the Tenant / Lessee shall abide by all the bye-laws, rules and regulations of the local authorities and shall not perform any illegal activities.",
        f"That this Lease is granted for a period of {data['rent'].get('duration', '11 months')} only commencing from {data['rent']['start_date']} and may be extended by mutual consent.",
        "That the Lessee shall pay electricity and water charges proportionately to the Lessor/Owner.",
        "That the Tenant/Lessee shall not make any structural changes without the consent of the Owner, except temporary decor like wooden partitions or AC units.",
        "That no addition or alteration is allowed without written consent, nor can the premises be sublet.",
        "That the Lessor/Owner or their agent may inspect the premises for repairs or checks with reasonable notice.",
        "That the Tenant/Lessee shall maintain cleanliness and hygiene and avoid any nuisance.",
        "That the Tenant/Lessee shall handle minor repairs at their own cost.",
        "That this Agreement can be terminated before expiry by giving one month's prior notice by either party.",
        "That the premises shall be used for official purposes only.",
        "That no dangerous, explosive, or unlawful items shall be stored on the premises.",
        "That the Lessee shall pay one month's advance rent which will be adjusted later.",
        "Both parties have read and understood the agreement and sign it without pressure."
    ]

    for clause in clauses:
        doc.add_paragraph(clause, style='List Number')

    doc.add_paragraph(
        f"\nIN WITNESS WHEREOF, the Lessor/Owner and the Tenant/Lessee have subscribed their hands at {data['place']} on {data['agreement_date']} in the presence of the following witnesses."
    )

    doc.add_paragraph("\nWITNESSES:\n1.\n2.")

    doc.add_paragraph(f"\n{data['landlord']['name']}\nLessor")
    doc.add_paragraph(f"\n{data['tenant']['company_name']}\nLessee")

    output_dir = "generated_docs"
    os.makedirs(output_dir, exist_ok=True)
    file_path = os.path.join(output_dir, f"{filename_prefix}.docx")
    doc.save(file_path)

    return file_path
