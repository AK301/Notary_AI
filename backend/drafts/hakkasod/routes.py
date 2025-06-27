from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from docx import Document
import os, uuid

from backend.drafts.hakkasod.generator import generate_hakkasod_docx
from backend.drafts.hakkasod.schema import HakkasodPatraRequest
from backend.drafts.hakkasod.generate_ai import (
    AIDraftRequest,
    generate_hakkasod_from_scenario,
)

router = APIRouter(prefix="/hakkasod", tags=["Hakkasod Patra Draft"])


@router.post("/generate")
def generate(data: HakkasodPatraRequest):
    path = generate_hakkasod_docx(data, filename_prefix=data.receiver_name)
    return {"docx_path": path}


@router.post("/generate-ai")
def generate_ai_draft(data: AIDraftRequest):
    content = generate_hakkasod_from_scenario(data)
    return {"draft": content}


@router.post("/generate-ai-docx")
def generate_ai_draft_and_docx(data: AIDraftRequest):
    try:
        content = generate_hakkasod_from_scenario(data)

        doc = Document()
        doc.add_heading("हक्कसोड पत्र", 0)
        doc.add_paragraph(content)

        filename = f"hakkasod_{uuid.uuid4().hex[:6]}.docx"
        folder = "generated_docs"
        os.makedirs(folder, exist_ok=True)
        path = os.path.join(folder, filename)
        doc.save(path)

        return FileResponse(
            path,
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        print("❌ Error in generate-ai-docx:", str(e))
        raise HTTPException(status_code=500, detail="Failed to generate document.")
