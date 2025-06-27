from docx import Document
from backend.drafts.hakkasod.schema import HakkasodPatraRequest


def generate_hakkasod_docx(data: HakkasodPatraRequest, filename_prefix: str):
    doc = Document()
    doc.add_heading("हक्कसोड पत्र", level=1)

    doc.add_paragraph(f"सन्माननीय अधिकारी,")
    doc.add_paragraph(f"{data.receiver_name} यांच्यासाठी खालील व्यक्ती हक्कसोड करीत आहेत:")

    for i, person in enumerate(data.relinquishers, start=1):
        doc.add_paragraph(f"{i}. {person.name}, {person.relation}, {person.occupation}")

    doc.add_paragraph("जमीन तपशील:")
    for land in data.land_details:
        doc.add_paragraph(f"- सर्वे नं {land.survey_number}, क्षेत्रफळ: {land.area}")

    doc.add_paragraph(f"गाव: {data.village}, जिल्हा: {data.district}")
    doc.add_paragraph(f"दिनांक: {data.date}")

    filename = f"{filename_prefix}_hakkasod.docx"
    filepath = f"./generated_docs/{filename}"
    doc.save(filepath)

    return filepath
