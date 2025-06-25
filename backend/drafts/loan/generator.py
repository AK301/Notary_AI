from docx import Document
import os
from datetime import datetime
from backend.drafts.loan.schema import LoanAgreementRequest


def generate_loan_docx(
    data: LoanAgreementRequest, filename_prefix: str = "loan_agreement"
) -> str:
    doc = Document()

    doc.add_heading("LOAN AGREEMENT", level=1)
    doc.add_paragraph("ACKNOWLEDGEMENT OF DEBT")

    doc.add_paragraph("Entered into between:")
    doc.add_paragraph(f'{data.lender_name}\n("The Lender")')
    doc.add_paragraph("and")
    doc.add_paragraph(f'{data.borrower_name}\n("The Borrower")')

    doc.add_paragraph("1. Amount of Loan")
    doc.add_paragraph(
        f"The Lender hereby agrees to lend the sum of ₹{data.loan_amount:,.2f} to the Borrower on the terms set out hereunder."
    )

    doc.add_paragraph("2. Payment of loan to Borrower")
    doc.add_paragraph(
        "It is agreed between the parties that payment of the loan amount will not be made to the Borrower before the expiry "
        "of three business days after the conclusion of the contract. During the said period, the Borrower may terminate the "
        "contract at will. It is further agreed that the Lender shall not be entitled to interest for the period preceding the date "
        "upon which the money is paid to the Borrower."
    )

    doc.add_paragraph("3. Period of Loan")
    doc.add_paragraph(
        f"This loan shall endure for a period of {data.loan_period_months} months calculated from {data.loan_start_date.strftime('%d-%m-%Y')}.\n"
        f"(In order to claim exemption from the Usury Act 73 of 1968 this number may not exceed 36 months.)"
    )

    doc.add_paragraph("4. Interest")
    if data.repayment_option.lower() == "emi" and data.monthly_installment_amount:
        doc.add_paragraph(
            f"The Borrower shall be obliged to pay interest at the rate of {data.interest_rate_percent:.2f}% per annum, "
            f"the interest and capital to be paid in equal monthly instalments of ₹{data.monthly_installment_amount:,.2f}."
        )
    else:
        doc.add_paragraph(
            f"The Borrower shall be obliged to pay interest at the rate of {data.interest_rate_percent:.2f}% per annum, "
            f"such interest to be paid together with the capital sum of the loan at the end of the loan period."
        )

    doc.add_paragraph("5. Exceptio non numeratae pecuniae")
    doc.add_paragraph(
        "The Borrower expressly renounces the benefit of the exception non numeratae pecuniae and confirms that the money has been received."
    )

    # Footer and signature section
    doc.add_paragraph(
        "\nIN WITNESS WHEREOF, the parties hereto have signed this agreement."
    )
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Lender Signature:\n\n\n__________________________"
    table.cell(0, 1).text = "Borrower Signature:\n\n\n__________________________"
    table.cell(1, 0).text = data.lender_name
    table.cell(1, 1).text = data.borrower_name

    output_dir = "generated_docs"
    os.makedirs(output_dir, exist_ok=True)
    filename = f"{filename_prefix.replace(' ', '_')}.docx"
    file_path = os.path.join(output_dir, filename)
    doc.save(file_path)

    return file_path
