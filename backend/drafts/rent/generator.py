from docx import Document
import os
from backend.drafts.rent.schema import RentAgreementRequest


def generate_rent_docx(
    data: RentAgreementRequest, filename_prefix: str = "rent_agreement"
) -> str:
    doc = Document()

    doc.add_heading("RENT AGREEMENT", level=1)
    doc.add_paragraph(
        f"This Rent Agreement is made and executed on this day between "
        f"{data.landlord_name}, residing at {data.landlord_address}, hereinafter called the 'Landlord' "
        f"and {data.tenant_name}, residing at {data.tenant_address}, hereinafter called the 'Tenant'."
    )

    doc.add_paragraph(
        f"The Landlord hereby agrees to let out the property located at {data.property_address} "
        f"to the Tenant at a monthly rent of ₹{data.rent_amount}, with a security deposit of ₹{data.deposit_amount}."
    )

    doc.add_paragraph(
        f"The rental term shall begin on {data.rent_start_date.strftime('%d-%m-%Y')} "
        f"and end on {data.rent_end_date.strftime('%d-%m-%Y')}."
    )

    doc.add_paragraph(
        "Both parties agree to abide by the terms and conditions set forth in this agreement."
    )

    doc.add_paragraph("\nIN WITNESS WHEREOF, the parties have signed this agreement.")
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Landlord Signature:\n\n\n__________________________"
    table.cell(0, 1).text = "Tenant Signature:\n\n\n__________________________"
    table.cell(1, 0).text = data.landlord_name
    table.cell(1, 1).text = data.tenant_name

    output_dir = "generated_docs"
    os.makedirs(output_dir, exist_ok=True)

    filename = f"{filename_prefix.replace(' ', '_')}.docx"
    file_path = os.path.join(output_dir, filename)
    doc.save(file_path)

    return file_path
